#!/usr/bin/env python3
from docx import Document
from pathlib import Path

docs = [
    '/Users/satish.korra/win-reports/Windows-Server-2019-CIS-Audit-Report.docx',
    '/Users/satish.korra/win-reports/All-2019/Windows-Server-2019-CIS-Audit-Report.docx',
    '/Users/satish.korra/win-reports/Prod-2019/Windows-Server-2019-CIS-Audit-Report.docx',
]

for doc_path in docs:
    doc_file = Path(doc_path)
    if not doc_file.exists():
        continue
    
    doc = Document(doc_file)
    
    for para in doc.paragraphs:
        text = para.text
        if 'Total Controls Evaluated: 332' in text:
            para.text = text.replace('Total Controls Evaluated: 332', 'Total Controls Evaluated: 334')
        if 'Passed: 88' in text:
            para.text = text.replace('Passed: 88', 'Passed: [TO BE UPDATED]')
        if 'Failed: 423' in text:
            para.text = text.replace('Failed: 423', 'Failed: [TO BE UPDATED]')
    
    # Check tables too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    if 'Total Controls Evaluated: 332' in text:
                        para.text = text.replace('Total Controls Evaluated: 332', 'Total Controls Evaluated: 334')
                    if 'Passed: 88' in text:
                        para.text = text.replace('Passed: 88', 'Passed: [TO BE UPDATED]')
                    if 'Failed: 423' in text:
                        para.text = text.replace('Failed: 423', 'Failed: [TO BE UPDATED]')
    
    doc.save(doc_file)
    print(f"✅ Updated {doc_file.name}")

print("\n✅ Executive summaries updated!")
print("   Total Controls: 334 (was 332)")
print("   Pass/Fail counts: Marked for update after scan")
