#!/usr/bin/env python3
"""Generate Windows Server 2019 CIS Audit Report (Word) with exclusions applied"""
import csv
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load exclusions
exclusions = set()
exclusion_file = Path('EXCLUSIONS_2019.txt')
if exclusion_file.exists():
    for line in exclusion_file.read_text().splitlines():
        line = line.strip()
        if line and not line.startswith('#') and '.' in line:
            exclusions.add(line)

print(f"Loaded {len(exclusions)} exclusions")

# Find CSV
csv_file = None
if len(sys.argv) > 1:
    csv_file = Path(sys.argv[1])
else:
    reports_dir = Path('reports')
    if reports_dir.exists():
        csv_files = list(reports_dir.glob('*.csv'))
        if csv_files:
            csv_file = sorted(csv_files, key=lambda x: x.stat().st_mtime)[-1]

if not csv_file or not csv_file.exists():
    print("❌ No CSV file found. Usage: python3 generate-word-report.py path/to/report.csv")
    sys.exit(1)

print(f"Processing: {csv_file}")

# Read CSV
controls = []
with open(csv_file, 'r', encoding='utf-8') as f:
    reader = csv.DictReader(f)
    for row in reader:
        if row.get('Control ID', '').strip() not in exclusions:
            controls.append(row)

failed = [c for c in controls if c.get('Status', '').lower() == 'fail']
passed = [c for c in controls if c.get('Status', '').lower() == 'pass']

# Create Word document
doc = Document()

# Title
title = doc.add_heading('Windows Server 2019 CIS Audit Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Summary
doc.add_heading('Executive Summary', 1)
p = doc.add_paragraph()
p.add_run(f"Total Controls Audited: ").bold = True
p.add_run(f"{len(controls)}\n")
p.add_run(f"Excluded (N/A): ").bold = True
p.add_run(f"{len(exclusions)}\n")
p.add_run(f"Passed: ").bold = True
p.add_run(f"{len(passed)}\n").font.color.rgb = RGBColor(0, 128, 0)
p.add_run(f"Failed: ").bold = True
p.add_run(f"{len(failed)}\n").font.color.rgb = RGBColor(255, 0, 0)
p.add_run(f"Compliance Rate: ").bold = True
p.add_run(f"{len(passed)/len(controls)*100:.1f}%\n")

# Failed controls
doc.add_heading('Failed Controls', 1)
doc.add_paragraph('The following controls require remediation:')

for control in failed:
    doc.add_heading(f"{control.get('Control ID', '')}: {control.get('Control Title', '')}", 2)
    doc.add_paragraph(f"Status: FAIL", style='List Bullet').runs[0].font.color.rgb = RGBColor(255, 0, 0)
    doc.add_paragraph('[Screenshot placeholder - Add screenshot here]', style='List Bullet')
    doc.add_paragraph()

# Passed controls
doc.add_page_break()
doc.add_heading('Passed Controls', 1)
for control in passed:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f"{control.get('Control ID', '')}: {control.get('Control Title', '')}").font.color.rgb = RGBColor(0, 128, 0)

# Excluded controls
doc.add_page_break()
doc.add_heading('Excluded Controls (Not Applicable)', 1)
doc.add_paragraph('The following controls are not applicable to Windows Server 2019:')

with open(exclusion_file, 'r') as f:
    for line in f:
        line = line.strip()
        if line and not line.startswith('#') and '.' in line:
            doc.add_paragraph(line, style='List Bullet')

# Save
output_file = Path('Windows-Server-2019-CIS-Audit-Report.docx')
doc.save(output_file)

print(f"\n✅ Report generated: {output_file}")
print(f"   Total controls: {len(controls)}")
print(f"   Failed: {len(failed)}")
print(f"   Passed: {len(passed)}")
print(f"   Excluded: {len(exclusions)}")
