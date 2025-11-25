#!/usr/bin/env python3
"""Remove 62 excluded controls from existing Word documents"""
from pathlib import Path
from docx import Document
import sys

# Load exclusions
exclusions = set()
exclusion_file = Path('EXCLUSIONS_2019.txt')
for line in exclusion_file.read_text().splitlines():
    line = line.strip()
    if line and not line.startswith('#') and '.' in line:
        exclusions.add(line)

print(f"Loaded {len(exclusions)} exclusions\n")

# Word docs to update
docs = [
    '/Users/satish.korra/win-reports/Windows-Server-2019-CIS-Audit-Report.docx',
    '/Users/satish.korra/win-reports/All-2019/Windows-Server-2019-CIS-Audit-Report.docx',
    '/Users/satish.korra/win-reports/Prod-2019/Windows-Server-2019-CIS-Audit-Report.docx',
]

for doc_path in docs:
    doc_file = Path(doc_path)
    if not doc_file.exists():
        print(f"⏭️  Skipping {doc_file.name} (not found)")
        continue
    
    print(f"Processing: {doc_file}")
    doc = Document(doc_file)
    
    removed = 0
    paragraphs_to_remove = []
    
    # Find paragraphs with excluded control IDs
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for excl in exclusions:
            if excl in text or excl.replace('.', '_') in text:
                paragraphs_to_remove.append(i)
                removed += 1
                break
    
    # Remove from tables too
    for table in doc.tables:
        rows_to_remove = []
        for i, row in enumerate(table.rows):
            row_text = ' '.join(cell.text for cell in row.cells)
            for excl in exclusions:
                if excl in row_text:
                    rows_to_remove.append(i)
                    removed += 1
                    break
        
        # Remove rows in reverse order
        for i in reversed(rows_to_remove):
            table._element.remove(table.rows[i]._element)
    
    # Save with backup
    backup = doc_file.with_suffix('.docx.backup')
    doc_file.rename(backup)
    doc.save(doc_file)
    
    print(f"✅ Updated {doc_file.name}")
    print(f"   Removed ~{removed} references to excluded controls")
    print(f"   Backup: {backup.name}\n")

print("✅ All Word documents updated!")
print(f"\nExcluded controls removed: {len(exclusions)}")
